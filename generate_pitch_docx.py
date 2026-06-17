"""
Generate a polished, creative PITCH ONE-PAGER in .docx format.
Focused on "What's in it for Eenadu/ETV" — investor-centric framing.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- Page margins (narrow for more content space) ---
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# --- Helper functions ---
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_paragraph(doc, text, size=11, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = 'Segoe UI'
    return p

def add_mixed_paragraph(doc, parts, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    """parts = list of (text, size, bold, color_tuple_or_None)"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for text, size, bold, color in parts:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = 'Segoe UI'
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def make_table_borderless(table):
    """Remove all borders from table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        borders.append(border)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)

# ============================================================
# HEADER — Bold brand statement
# ============================================================

# Title
add_styled_paragraph(doc, "KSHETRA", size=28, bold=True, color=(26, 26, 26), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_styled_paragraph(doc, "India's Political Intelligence Platform", size=14, bold=False, color=(100, 100, 100), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

# Tagline / Hook
add_styled_paragraph(doc, "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━", size=8, color=(200, 200, 200), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

add_mixed_paragraph(doc, [
    ("THE OPPORTUNITY: ", 11, True, (180, 60, 20)),
    ("India's delimitation — the first redrawing of constituency boundaries in 50+ years — creates a ", 11, False, (50, 50, 50)),
    ("once-in-a-generation media opportunity", 11, True, (50, 50, 50)),
    (". The media house that owns the data wins the decade.", 11, False, (50, 50, 50)),
], alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)

# ============================================================
# SECTION: WHAT'S IN IT FOR EENADU/ETV (Lead with THEIR benefit)
# ============================================================

add_styled_paragraph(doc, "WHAT EENADU/ETV GETS", size=14, bold=True, color=(26, 26, 26), space_after=8)

# Benefits table
benefits = [
    ("🏆  Election Night Dominance", 
     "Real-time constituency maps on ETV screens. While competitors show static graphics, you show LIVE interactive data for 4,000+ seats. Viewers stay glued."),
    ("📱  A Ready-Made Digital Product",
     "No 12-month development cycle. No ₹2-3 Cr tech investment. The app is BUILT. Launch under 'ETV Election Command' branding within weeks."),
    ("🎯  Constituency-Level Ad Targeting",
     "Sell political ads at ₹5-10x premium because candidates can target ONLY their constituency. No other platform in India offers this."),
    ("📊  Newsroom Intelligence",
     "Your journalists get instant access to MLA assets, criminal records, vote margins, swing analysis. Faster stories. Better scoops."),
    ("🗺️  Delimitation Exclusive",
     "India's ONLY working delimitation simulator. First media house to explain \"how your constituency changes\" wins massive trust and viewership."),
    ("🌐  Multi-State = Multi-Channel",
     "Works for ETV Telugu + ETV Bangla + ETV Kannada + ETV Marathi + all other language channels. One product powers your entire network."),
]

for title, desc in benefits:
    add_mixed_paragraph(doc, [
        (title, 11, True, (26, 26, 26)),
    ], space_after=2)
    add_styled_paragraph(doc, desc, size=10, color=(80, 80, 80), space_after=10)

# ============================================================
# SECTION: WHAT'S BUILT (proof it's real)
# ============================================================

add_styled_paragraph(doc, "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━", size=8, color=(200, 200, 200), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

add_mixed_paragraph(doc, [
    ("THIS IS NOT A CONCEPT. ", 12, True, (180, 60, 20)),
    ("IT'S BUILT AND WORKING.", 12, True, (180, 60, 20)),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# Feature grid as table
table = doc.add_table(rows=5, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

features_col1 = [
    ("✅ 31 States Covered", "4,000+ constituencies mapped"),
    ("✅ MLA/MP X-Ray", "Assets, criminal cases, wealth growth"),
    ("✅ Delimitation Simulator", "India's ONLY — predict new boundaries"),
    ("✅ Election Analytics", "Swing seats, anti-incumbency, heatmaps"),
    ("✅ Promise Tracker", "Govt promises: kept, broken, pending"),
]

features_col2 = [
    ("✅ Real-Time Political Timeline", "Defections, splits, by-elections"),
    ("✅ Civic Dashboard", "Issues, sentiment, MLA responsiveness"),
    ("✅ 5 Languages", "EN, HI, TE, KN, MR (more planned)"),
    ("✅ Offline-First", "Works without internet"),
    ("✅ Multi-Platform Ready", "Android live, iOS & Web in pipeline"),
]

for i in range(5):
    cell1 = table.cell(i, 0)
    cell2 = table.cell(i, 1)
    
    # Column 1
    p1 = cell1.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    run1 = p1.add_run(features_col1[i][0])
    run1.font.size = Pt(10)
    run1.font.bold = True
    run1.font.name = 'Segoe UI'
    p1_sub = cell1.add_paragraph()
    p1_sub.paragraph_format.space_after = Pt(6)
    run1_sub = p1_sub.add_run(features_col1[i][1])
    run1_sub.font.size = Pt(9)
    run1_sub.font.color.rgb = RGBColor(100, 100, 100)
    run1_sub.font.name = 'Segoe UI'
    
    # Column 2
    p2 = cell2.paragraphs[0]
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run(features_col2[i][0])
    run2.font.size = Pt(10)
    run2.font.bold = True
    run2.font.name = 'Segoe UI'
    p2_sub = cell2.add_paragraph()
    p2_sub.paragraph_format.space_after = Pt(6)
    run2_sub = p2_sub.add_run(features_col2[i][1])
    run2_sub.font.size = Pt(9)
    run2_sub.font.color.rgb = RGBColor(100, 100, 100)
    run2_sub.font.name = 'Segoe UI'

make_table_borderless(table)

# ============================================================
# SECTION: THE NUMBERS THAT MATTER
# ============================================================

doc.add_paragraph()  # spacing
add_styled_paragraph(doc, "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━", size=8, color=(200, 200, 200), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_styled_paragraph(doc, "THE NUMBERS THAT MATTER", size=14, bold=True, color=(26, 26, 26), space_after=10)

# Key metrics in a visual way
metrics_table = doc.add_table(rows=1, cols=4)
metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER

metrics = [
    ("900M+", "Voters in India"),
    ("₹15,000 Cr", "Election Ad Spend/Cycle"),
    ("ZERO", "Competitors at this depth"),
    ("50+ Years", "Since last delimitation"),
]

for i, (number, label) in enumerate(metrics):
    cell = metrics_table.cell(0, i)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(number)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(180, 60, 20)
    run.font.name = 'Segoe UI'
    
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(label)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(100, 100, 100)
    run2.font.name = 'Segoe UI'

make_table_borderless(metrics_table)

# ============================================================
# SECTION: WHY THIS, WHY NOW, WHY YOU
# ============================================================

doc.add_paragraph()
add_styled_paragraph(doc, "WHY EENADU/ETV — WHY NOW", size=14, bold=True, color=(26, 26, 26), space_after=8)

why_points = [
    ("Delimitation is coming.", "The first media house to explain it wins the narrative. You can be that house — starting tomorrow."),
    ("Your network is the distribution.", "ETV's 10+ language channels reach the exact audience this product serves. No marketing spend needed — just integrate."),
    ("Election cycles are perpetual.", "Every 6 months, some state votes. This isn't a one-time product — it's a permanent competitive advantage."),
    ("Build vs. Buy.", "Building this in-house would cost ₹2-3 Cr and 12+ months. Or you can have it now, proven, for a fraction."),
]

for title, desc in why_points:
    add_mixed_paragraph(doc, [
        (f"▸ {title} ", 11, True, (26, 26, 26)),
        (desc, 10, False, (80, 80, 80)),
    ], space_after=8)

# ============================================================
# SECTION: THE ASK (clean, simple)
# ============================================================

add_styled_paragraph(doc, "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━", size=8, color=(200, 200, 200), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_styled_paragraph(doc, "THE PARTNERSHIP", size=14, bold=True, color=(26, 26, 26), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# Two-column: What You Get / What We Need
ask_table = doc.add_table(rows=1, cols=2)
ask_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Left: What they get
left_cell = ask_table.cell(0, 0)
p = left_cell.paragraphs[0]
run = p.add_run("WHAT YOU GET")
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = RGBColor(34, 139, 34)
run.font.name = 'Segoe UI'

left_items = [
    "Exclusive media licensing (AP/TS or national)",
    "White-label 'ETV Election Command' branding",
    "Equity stake in a ₹500Cr+ potential company",
    "First access to every new feature",
    "Constituency-level ad inventory (new revenue)",
    "Newsroom intelligence for all journalists",
]
for item in left_items:
    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  ✓  {item}")
    run.font.size = Pt(9.5)
    run.font.name = 'Segoe UI'

# Right: What we need
right_cell = ask_table.cell(0, 1)
p = right_cell.paragraphs[0]
run = p.add_run("WHAT WE NEED")
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = RGBColor(180, 60, 20)
run.font.name = 'Segoe UI'

right_items = [
    "Strategic investment (₹75L – 1 Cr)",
    "Distribution via ETV/Eenadu audience",
    "Content team support (election data entry)",
    "Office/infrastructure for small team",
    "Introductions to national media partners",
    "6-month exclusivity window",
]
for item in right_items:
    p = right_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  ▸  {item}")
    run.font.size = Pt(9.5)
    run.font.name = 'Segoe UI'

make_table_borderless(ask_table)

# ============================================================
# FOOTER
# ============================================================

doc.add_paragraph()
add_styled_paragraph(doc, "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━", size=8, color=(200, 200, 200), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

add_mixed_paragraph(doc, [
    ("\"The media house that owns India's political data ", 11, False, (100, 100, 100)),
    ("owns the election narrative for the next decade.\"", 11, True, (26, 26, 26)),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_styled_paragraph(doc, "[Your Name]  |  Founder, Kshetra  |  [Your Phone]  |  Hyderabad", size=10, color=(100, 100, 100), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_styled_paragraph(doc, "Detailed 5-year financial projections available as Annexure on request.", size=9, color=(150, 150, 150), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

# ============================================================
# SAVE
# ============================================================

output_path = r"C:\Users\Laven\OneDrive\Desktop\Kshetra\Kshetra_Pitch_OnePager.docx"
doc.save(output_path)
print(f"✅ Saved: {output_path}")
print(f"   Size: {os.path.getsize(output_path) / 1024:.1f} KB")
