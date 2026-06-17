"""
Generate a non-technical-friendly TECH STACK document for KSHETRA in .docx format.
Audience: a non-technical founder presenting to an investor's project vetting team.
Every technology is explained in plain English with a simple analogy.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# --- Brand colors ---
NAVY = (10, 10, 26)        # #0A0A1A
SAFFRON = (255, 153, 51)   # India saffron accent
GREEN = (19, 136, 8)
GREY = (90, 90, 90)
WHITE = (255, 255, 255)
LIGHT = (245, 246, 250)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def rgb_hex(color):
    return '%02X%02X%02X' % color


def para(text, size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_after=6, space_before=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = 'Segoe UI'
    return p


def heading(text, size=15):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*NAVY)
    run.font.name = 'Segoe UI'
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), rgb_hex(SAFFRON))
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(10.5)
        r.font.name = 'Segoe UI'
        r.font.color.rgb = RGBColor(*NAVY)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = 'Segoe UI'
    return p


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], rgb_hex(NAVY))
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*WHITE)
        run.font.name = 'Segoe UI'
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if r_idx % 2 == 1:
                set_cell_shading(cells[i], rgb_hex(LIGHT))
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            run.font.name = 'Segoe UI'
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(*NAVY)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ============================================================
# TITLE BLOCK
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
r = title.add_run('KSHETRA')
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = RGBColor(*NAVY)
r.font.name = 'Segoe UI'

para('Technology Stack — Explained in Plain English',
     size=13, bold=True, color=SAFFRON, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para('A guide for the investor project-vetting team',
     size=10.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=2)
para('Prepared for non-technical readers  |  Version 1.0',
     size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# ============================================================
# HOW TO USE THIS DOCUMENT
# ============================================================
heading('How to read this document')
para('This document lists every major technology used to build KSHETRA and explains each one in '
     'simple, everyday language. For each item you will see WHAT it is, WHY we chose it, and a SIMPLE '
     'ANALOGY you can use to explain it to anyone. You do not need a technical background to present this.',
     size=10.5, space_after=6)

# ============================================================
# 30-SECOND SUMMARY
# ============================================================
heading('The 30-second summary (say this first)')
para('"KSHETRA is built like a modern, professional tech product. We use the same proven, industry-standard '
     'tools that companies like Instagram, Discord and Shopify use. One single codebase powers our Android app, '
     'iPhone app, and website, which keeps our team small and our costs low. Our data lives in a secure, '
     'globally trusted cloud database, and we add Artificial Intelligence on top to give users smart insights."',
     size=10.5, italic=True, space_after=6)

para('The product is organised into four simple parts:', size=10.5, bold=True, space_after=4)
bullet('The phone app that citizens, candidates and journalists actually use.', '1. The App (front-end): ')
bullet('The "engine room" that handles requests, security and business logic.', '2. The Server (back-end): ')
bullet('The secure vault where all information is stored and searched.', '3. The Database: ')
bullet('The smart layer that answers questions and analyses political data.', '4. The AI brain: ')

# ============================================================
# SECTION 1: THE APP
# ============================================================
heading('Part 1 — The Mobile App (what users see and touch)')
make_table(
    ['Technology', 'What it is (plain English)', 'Why it matters'],
    [
        ['React Native + Expo',
         'A toolkit that lets us build the Android app, iPhone app, and website from ONE codebase instead of three.',
         'Cuts development cost and time by roughly two-thirds; one team maintains everything.'],
        ['TypeScript',
         'The programming language we write in. It is a safer version of JavaScript that catches mistakes before users ever see them.',
         'Fewer bugs, easier for new engineers to join and understand the code.'],
        ['Expo Router',
         'The system that controls moving between screens (e.g. Home to Search to Profile).',
         'Smooth, app-store-quality navigation that feels familiar to users.'],
        ['MapLibre (maps)',
         'An open-source mapping engine that draws interactive constituency maps and boundaries.',
         'No expensive per-map licence fees; full control over our political map data.'],
        ['Zustand + React Query',
         'The app\'s short-term memory. It keeps track of what the user is doing and fetches fresh data efficiently.',
         'Fast, responsive screens and less mobile data usage.'],
        ['FlashList + Reanimated',
         'Tools that make long lists scroll smoothly and animations feel fluid.',
         'A premium, polished feel that builds user trust.'],
        ['i18next (languages)',
         'The translation system that lets the app speak multiple Indian languages.',
         'Essential for nationwide reach across India\'s many languages.'],
    ],
    col_widths=[3.5, 7.5, 5.5]
)

# ============================================================
# SECTION 2: THE SERVER
# ============================================================
heading('Part 2 — The Server / Back-end (the engine room)')
make_table(
    ['Technology', 'What it is (plain English)', 'Why it matters'],
    [
        ['Node.js + Fastify',
         'The server software that receives requests from the app and sends back the right answers, very quickly.',
         'Fastify is one of the fastest options available; handles many users with low server cost.'],
        ['TypeScript',
         'Same safe language as the app, used on the server too.',
         'One language across the whole project = a smaller, more efficient team.'],
        ['Zod (data checks)',
         'A gatekeeper that verifies incoming information is valid and safe before it is used.',
         'Protects against bad or malicious data — a key security and reliability layer.'],
        ['Helmet + Rate-Limit + CORS',
         'A set of standard security guards that block common web attacks and abuse.',
         'Industry-standard protection that vetting teams expect to see.'],
    ],
    col_widths=[3.5, 7.5, 5.5]
)

# ============================================================
# SECTION 3: DATABASE
# ============================================================
heading('Part 3 — The Database (the secure vault)')
make_table(
    ['Technology', 'What it is (plain English)', 'Why it matters'],
    [
        ['Supabase',
         'A managed cloud platform that hosts our database, user logins, and file storage — all in one.',
         'Trusted by thousands of companies; gives us enterprise features without a large infrastructure team.'],
        ['PostgreSQL',
         'The actual database — the world\'s most respected open-source system for storing structured data.',
         'Rock-solid, proven at massive scale, no licence fees.'],
        ['PostGIS (maps data)',
         'An add-on that lets the database understand geography — "which constituency is this point inside?"',
         'Powers our core map and constituency features.'],
        ['Secure login & storage',
         'Built-in user sign-in, permissions, and secure storage of photos/evidence uploaded by users.',
         'Keeps citizen data private and access-controlled.'],
    ],
    col_widths=[3.5, 7.5, 5.5]
)

# ============================================================
# SECTION 4: AI
# ============================================================
heading('Part 4 — The AI Brain (smart insights)')
make_table(
    ['Technology', 'What it is (plain English)', 'Why it matters'],
    [
        ['OpenAI',
         'The artificial-intelligence service that powers our in-app assistant and analysis features.',
         'Lets users ask questions in plain language and get instant, intelligent answers.'],
        ['Content moderation',
         'Automated checks that screen user-submitted content for safety before it goes live.',
         'Critical for a civic/political platform to stay trustworthy and compliant.'],
    ],
    col_widths=[3.5, 7.5, 5.5]
)

# ============================================================
# SECTION 5: HOW IT IS BUILT & SHIPPED
# ============================================================
heading('Part 5 — How the project is organised and shipped')
make_table(
    ['Technology', 'What it is (plain English)', 'Why it matters'],
    [
        ['Turborepo (monorepo)',
         'A way of keeping the app, server, and shared code together in one tidy, well-organised repository.',
         'Code is reused, not duplicated; everything stays consistent.'],
        ['Shared package',
         'A common toolbox of rules and data types used by both the app and the server.',
         'Guarantees the app and server always "speak the same language".'],
        ['EAS Build (Expo)',
         'The cloud service that turns our code into the actual installable Android/iPhone app.',
         'Professional, repeatable app builds without owning expensive build machines.'],
        ['GitHub Actions',
         'Robots that automatically test the code every time we make a change.',
         'Catches problems early; a sign of a mature engineering process.'],
        ['Jest (testing)',
         'The automated testing framework that checks features still work after changes.',
         'Quality assurance built into the workflow.'],
        ['Python scrapers',
         'Separate data-collection scripts that gather and prepare political/election data.',
         'Feeds the platform with accurate, up-to-date constituency information.'],
    ],
    col_widths=[3.5, 7.5, 5.5]
)

# ============================================================
# WHY THIS STACK IS A STRENGTH
# ============================================================
heading('Why this technology choice is a strength (for the investor)')
bullet('Every tool is a widely-adopted industry standard — low risk, easy to hire for, and well-documented.', 'Proven & mainstream: ')
bullet('One codebase serves Android, iPhone and web, so a small team can do the work of three.', 'Cost-efficient: ')
bullet('Open-source foundations (PostgreSQL, MapLibre, Node.js) mean minimal licence fees as we grow.', 'Low ongoing cost: ')
bullet('The architecture is designed to handle growth from thousands to millions of users.', 'Scalable: ')
bullet('Security guards, data validation, and content moderation are built in from day one.', 'Secure by design: ')

# ============================================================
# GLOSSARY
# ============================================================
heading('Mini-glossary (quick definitions)')
glossary = [
    ('Codebase', 'All the written code that makes up the product.'),
    ('Front-end', 'The part users see and interact with (the app screens).'),
    ('Back-end', 'The behind-the-scenes server that does the heavy lifting.'),
    ('Database', 'The organised storage where all information lives.'),
    ('API', 'The messenger that lets the app and server talk to each other.'),
    ('Open-source', 'Free, community-maintained software with no licence fee.'),
    ('Cloud', 'Computers run by a provider over the internet, rented as needed.'),
    ('Monorepo', 'One single, organised home for all the project\'s code.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(term + ' — ')
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(*NAVY)
    r.font.name = 'Segoe UI'
    r2 = p.add_run(definition)
    r2.font.size = Pt(10)
    r2.font.name = 'Segoe UI'

para('', size=6, space_after=4)
para('Note: This document describes the technologies actually in use in the current codebase. '
     'Some tools may evolve as the product matures.',
     size=8.5, color=GREY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# SAVE
# ============================================================
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'KSHETRA_Tech_Stack_Explained.docx')
doc.save(out_path)
print('Saved:', out_path)
